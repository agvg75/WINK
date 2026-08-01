import json
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

VERSION = json.loads(
    (Path(__file__).parent / "staged" / "app" / "release_info.json").read_text()
)["app_version"]
OUT = Path(__file__).parent / "staged" / "docs" / f"WINK_Lab_Tools_Definitive_Manual_v{VERSION}.docx"
NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
MAROON = RGBColor(153, 0, 0)
GRAY = RGBColor(90, 90, 90)
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = RGBColor(255, 255, 255)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(.492)
    section.footer_distance = Inches(.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(.375)
        style.paragraph_format.first_line_indent = Inches(-.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("AGVG LAB TOOLS  |  BIOLOGY-FIRST USER AND METHODS MANUAL")
    set_font(r, size=8.5, color=GRAY, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Page ")
    set_font(r, size=8.5, color=GRAY)
    add_field(footer, "PAGE")


def title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("WINK - Worm Imaging and Kinematics  (formerly NIKE)")
    set_font(r, size=12, color=MAROON, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Biology-First User and Methods Manual")
    set_font(r, size=28, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("How to choose, operate, tune, review, and interpret the Vidal-Gadea Lab analysis ecosystem")
    set_font(r, size=13.5, color=GRAY, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(58)
    r = p.add_run("A living manual for C. elegans behavior, physiology, and morphology")
    set_font(r, size=11, color=MAROON, bold=True)
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2700, 6660])
    items = [
        ("Audience", "Worm scientists, students, and collaborators using the Lab Hub"),
        ("Manual version", "1.1 (WINK rename)"),
        ("Software snapshot", f"WINK Lab Hub application v{VERSION}; supervised segmentation, human-in-the-loop review, prototype learning, camera compensation, and temporal identity continuity"),
        ("Prepared", date.today().strftime("%B %d, %Y")),
    ]
    for row, (label, value) in zip(table.rows, items):
        shade(row.cells[0], LIGHT)
        set_font(row.cells[0].paragraphs[0].add_run(label), size=9.5, color=NAVY, bold=True)
        set_font(row.cells[1].paragraphs[0].add_run(value), size=9.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Design principle: begin with the biological question, expose assumptions that vary with biology, and preserve evidence for human review.")
    set_font(r, size=10.5, color=GRAY, italic=True)
    doc.add_page_break()


def add_callout(doc, label, text, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run(label + "  ")
    set_font(r, size=10, color=MAROON, bold=True)
    r = p.add_run(text)
    set_font(r, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(header), size=9, color=NAVY, bold=True)
    for record in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, record):
            p = cell.paragraphs[0]
            set_font(p.add_run(str(value)), size=9)
    return table


def add_labeled(doc, label, text):
    p = doc.add_paragraph()
    keep_with_next(p)
    r = p.add_run(label + ": ")
    set_font(r, color=NAVY, bold=True)
    p.add_run(text)


def tool(doc, name, status, aim, use, decisions, controls, mechanism, outputs, failures, special):
    doc.add_heading(name, level=2)
    add_labeled(doc, "Status", status)
    add_labeled(doc, "Biological aim", aim)
    add_labeled(doc, "Use it when", use)
    doc.add_heading("Biology-driven design decisions", level=3)
    add_bullets(doc, decisions)
    doc.add_heading("User-adjustable variables and why they exist", level=3)
    add_matrix(doc, ["Variable or choice", "Why it is exposed"], controls, [3000, 6360])
    doc.add_heading("How the tool accomplishes the measurement", level=3)
    for paragraph in mechanism:
        doc.add_paragraph(paragraph)
    doc.add_heading("Outputs and interpretation", level=3)
    add_bullets(doc, outputs)
    doc.add_heading("Known failure modes and cautions", level=3)
    add_bullets(doc, failures)
    doc.add_heading("Special-case tuning", level=3)
    add_bullets(doc, special)


def before_you_film(doc):
    doc.add_heading("Before you film: acquisition planning for WINK modules", level=2)
    doc.add_paragraph(
        "Most WINK modules are not magic filters applied after the fact. They are human-in-the-loop "
        "measurement assistants. They work best when the biological feature is plainly visible to a trained "
        "person and the recording has enough spatial and temporal information for that feature to be measured. "
        "A good movie reduces QC time; a borderline movie turns the user into the algorithm's life-support system.")
    add_callout(
        doc, "The practical rule",
        "Before collecting a full experiment, record a short pilot movie and open it in the intended module. "
        "If the feature, scale, frame rate, or field of view is already questionable in the pilot, fix acquisition "
        "before making a day of movies that all need rescue.")
    add_bullets(doc, [
        "Lock exposure, gain, illumination, and focus when possible. Auto-exposure and flicker can look like biology to a detector.",
        "Record frame rate, exposure, objective/magnification, scale calibration, channel order, temperature, strain, stage, and assay geometry with the raw data.",
        "Use enough zoom to resolve the structure being measured, but enough field of view to keep the animal, event, stimulus, or arena visible.",
        "Avoid crowding when identity or one anatomical event matters. Crowding is acceptable only when the module is explicitly population-level and has review tools for collisions.",
        "Keep raw files. Avoid lossy compression for fluorescence, morphology, faint neurons, eggs, or low-contrast DIC features.",
        "If a movie contains camera bumps, focus drift, out-of-frame intervals, larvae, debris, overlapping animals, or uneven illumination, plan extra review time and collect extra biological replicates.",
    ])
    doc.add_heading("Module-by-module filming guidance", level=3)
    add_matrix(doc, ["Module or biological question", "Film to make this easy", "Known hard cases / plan extra review"], [
        ("Track one worm / single-worm kinematics",
         "One isolated animal, fully in frame, with the whole body visible at a frame rate appropriate for crawl, swim, or burrow.",
         "Coils, self-contact, leaving the field, low body contrast, abrupt focus/exposure changes, and crowded animals."),
        ("Population tracking",
         "Whole useful arena visible; enough contrast for worms; resolution and frame rate matched to the behavior rather than maximal camera settings.",
         "Very high-resolution movies, dense collisions, long track gaps, worms at the edge, and ROIs that were not checked across the movie."),
        ("Basal slowing",
         "Stable view of the bacterial lawn boundary or food region with visible worms before and after entry; record timing and geometry.",
         "Camera motion, ambiguous food boundaries, worms hidden under clumps, missing FPS/scale, and recordings that start after the key transition."),
        ("Orientation / chemotaxis / magnetotaxis / thermotaxis",
         "Include the full assay geometry and document stimulus direction, temperature, magnet position, or gradient map.",
         "Stimulus not visible or not recorded, crowded plate edges, uneven illumination, and treating many worms on one plate as independent plates."),
        ("Pharyngeal pumping",
         "Zoom on the head/pharynx with stable focus, unsaturated contrast, and frame rate high enough to resolve individual pumps.",
         "Crawling heads that leave the ROI, motion blur, low grinder contrast, saturated brightfield edges, and too-low FPS."),
        ("Defecation / pBoc",
         "Best case: one adult worm, zoomed enough to see posterior motion, fairly stationary, isolated, and recorded long enough for multiple cycles.",
         "Nightmare case: crawling or turning worm on a crowded plate, larvae/debris near the tail, coils, tail hidden, or field too wide to see axial tail motion."),
        ("Endpoint egg counting",
         "Static or nearly static image, eggs in focus, scale recorded, and illumination that gives eggs a consistent oval edge/contrast signature.",
         "Debris the same size as eggs, overlapping eggs, eggs clipped by frame edges, strong agar texture, uneven illumination, and overexposed worms."),
        ("Dynamic egg laying",
         "A field that keeps the worm and newly laid eggs visible over time; stable camera; enough time before and after laying events.",
         "Worm body obscuring eggs, moving debris, eggs leaving the field, and recordings where events are visible only after heavy contrast stretching."),
        ("RGBCaMP / calcium in moving worms",
         "Raw, unsaturated fluorescence channels with visible body outline and stable channel order; enough spatial detail for ROI review.",
         "Bleed-through, saturation, lossy compression, body coils, neuron/ROI leaving the frame, and masks that cannot be inspected later."),
        ("Neuron tracker",
         "Neuron bright enough to see by eye, body outline visible, worm mostly in frame, and frame rate matched to the biological response.",
         "Low contrast that makes jagged spines, neuron leaving frame, excessive resolution without ROI/virtual-stack planning, and overlays hiding correction landmarks."),
        ("Nonstriated muscle morphology",
         "High-quality still or stack with scale, orientation, tissue boundaries, attachment sites, and unsaturated signal visible.",
         "Saturated puncta, cropped anatomy, missing dorsal/ventral or anterior/posterior orientation, high background, and all-four-muscle ambiguity."),
        ("Evoked mechanosensation",
         "Video or image sequence that shows both worm and stimulus timing/location; include baseline and recovery.",
         "Stimulus present only in a separate note/CSV, worm not visible during stimulus, or no frame-synchronized event marker."),
    ], [2450, 3450, 3460])
    doc.add_heading("Visual examples to collect with each module", level=3)
    doc.add_paragraph(
        "For every module, the lab should maintain a small visual library: one good raw frame, one good overlay "
        "showing the detected feature, and one or two examples where the tool struggled. These images train users "
        "to recognize acquisition conditions before they lose time to QC. They also make the analysis easier to "
        "defend in peer review because the manual shows what the module is actually trying to identify.")
    add_matrix(doc, ["Example type", "What to show", "Why it helps"], [
        ("Good raw frame", "The biological feature before any overlay: worm, pharynx, tail, egg, neuron, muscle, or stimulus.",
         "Teaches users what the movie should look like before analysis."),
        ("Successful detection overlay", "The same frame with mask, spine, ROI, event marker, vector, or candidate egg labels.",
         "Shows what the module means by a correct detection."),
        ("Difficult-but-salvageable frame", "A low-contrast, crowded, partial, or unevenly illuminated example that still can be reviewed.",
         "Helps users plan extra review time rather than discard usable data prematurely."),
        ("Bad acquisition example", "A movie/frame where the relevant biology is not visible or is confounded beyond reliable correction.",
         "Gives students permission to reacquire instead of tuning forever."),
    ], [2000, 4100, 3260])
    add_callout(
        doc, "Human training matters too",
        "The goal is not to replace the user. The goal is to help students learn what the module is looking for, "
        "then let WINK propose, measure, save provenance, and make correction faster.")
    doc.add_heading("Minimum acquisition note to save with raw data", level=3)
    add_bullets(doc, [
        "Module intended and biological question.",
        "Strain, stage/age, treatment, temperature, assay plate, and biological replicate ID.",
        "Objective/magnification, camera, FPS, exposure, gain, bit depth, compression, and recording duration.",
        "Scale calibration method or known micrometers per pixel.",
        "Channel order and fluorophores, if applicable.",
        "Stimulus timing, direction, intensity, and geometry, if applicable.",
        "Known artifacts: camera movement, focus drift, crowded animals, larvae, debris, plate edge, clipping, or illumination problems.",
    ])


def front_matter(doc):
    doc.add_heading("How to use this manual", level=1)
    doc.add_paragraph(
        "Part I gets a new user from raw recording to reviewed result. Part II explains the "
        "shared scientific design of the ecosystem. Parts III-VII document every current Hub "
        "entry by category. Read the complete chapter for any tool before using its outputs as "
        "publication-grade measurements.")
    add_callout(doc, "Living-document rule",
                "When a tool changes, update its aim, exposed controls, measurement definition, "
                "outputs, validation status, and failure modes in the same software release.")
    doc.add_heading("Contents", level=1)
    for text in (
        "Part I - Get running quickly",
        "Part II - Shared scientific and computational principles",
        "Part III - Motor output: locomotion and sensory-guided behavior",
        "Part IV - Motor output: rhythmic programs",
        "Part V - Physiology: calcium and cellular activity",
        "Part VI - Anatomy and morphology",
        "Part VII - Mechanosensation and evoked response",
        "Part VIII - Acquisition and utilities",
        "Part IX - Validation, statistics, troubleshooting, and maintenance",
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)

    doc.add_heading("Part I - Get running quickly", level=1)
    doc.add_heading("First use on a lab computer", level=2)
    add_numbered(doc, [
        "Run Install_Lab_Tools.bat once on the computer.",
        "Launch WINK Lab Tools from the desktop shortcut.",
        "Choose a category in the left rail, or search by tool, measurement, requirement, or category.",
        "Select a tool card and read What it measures and Needs at load in the docked detail pane.",
        "Use the Ready or Experimental filter when validation status matters. External identifies third-party software outside the WINK validation ladder.",
        "Load a movie, stack, or folder when a movie tool should receive it directly, then choose Launch in the detail pane.",
        "Probe unfamiliar movies before analysis; convert them when Fiji or a tracker needs a clean TIFF stack.",
        "Enter acquisition metadata from the recording, never from a guess.",
        "Review every flagged track, event, ROI, interval, or segmentation before accepting the result.",
        "Keep raw data, reviewed outputs, metadata, ROI files, and the exact installer version together.",
    ])
    add_callout(
        doc, "Status chips are validation information",
        "Ready means technical validation or higher. Experimental means "
        "computational regression. External means another project defines "
        "the outputs and WINK does not apply its validation stamp.")
    doc.add_heading("Launcher controls", level=2)
    add_bullets(doc, [
        "Category rail: shows every registry category and its current tool count.",
        "Tool cards: show the registry name, one-line description, and validation-derived status word.",
        "Live filter: searches all categories and labels every result with its category.",
        "All tools: jumps directly to any registry entry and opens its details.",
        "Status filter: restricts cards to All, Ready, or Experimental.",
        "Keyboard: Tab reaches controls and cards; arrows move through categories or cards; Enter selects a card; the detail Launch button starts it.",
        "Check for updates: a local installation applies a checksum-verified atomic update. A versioned shared L-drive release opens the separately published newer folder and never renames or deletes shared files.",
    ])
    add_callout(
        doc, "Shared-folder update rule",
        "Do not grant students write access merely to update WINK The lab publishes a new versioned Current_Files folder; the running Hub validates that release stamp, opens it with the existing runtime, and closes the older Hub.")
    doc.add_heading("Tierpsy Tracker (External)", level=2)
    doc.add_paragraph(
        "Tierpsy is not bundled. Install it independently and set TIERPSY_PATH "
        "to its executable. The External card becomes launchable when the path "
        "is found; otherwise it says Not configured.")
    add_callout(
        doc, "Do not merge definitions silently",
        "Tierpsy WCON and feature files are Tierpsy-defined. They do not carry "
        "WINK QC, validation stamps, or plate-as-replicate rules. Do not "
        "compare Tierpsy and WINK outputs as though identically named "
        "features have identical definitions.")
    doc.add_heading("The five questions to answer before choosing a tool", level=2)
    add_bullets(doc, [
        "What is the biological unit: one worm, a population, one plate, one muscle, one neuron, or one egg?",
        "Is identity through time required, or is an identity-free population state more defensible?",
        "Is the desired quantity a posture, event, rate, trajectory, spatial distribution, fluorescence signal, or morphology?",
        "Which acquisition properties are known: frame rate, exposure, spatial calibration, bit depth, channel identity, and orientation?",
        "Which decisions require human biological judgment rather than a fixed algorithm?",
    ])
    doc.add_heading("Fast tool-selection guide", level=2)
    add_matrix(doc, ["Question", "Start with", "Do not substitute"], [
        ("One worm's signed bending and movement", "Track one worm; then Kinematics analysis/browser", "Population centroid tracking"),
        ("One worm's swimming bouts", "Single-worm swimming analysis", "Population modality proposals without review"),
        ("Many worms swimming/crawling/burrowing", "Population tracking", "One frequency threshold alone"),
        ("Response to OP50 entry", "Population basal slowing", "Unpaired population averages"),
        ("Plate orientation without reliable identity", "Population orientation; then Combine orientation plates", "Treating worms as independent N"),
        ("Pumps, pBoc, or eggs", "The matching rhythmic-program tool", "Generic motion peaks without event review"),
        ("Muscle calcium in moving animals", "RGBCaMP extractor + analysis/browser", "Raw intensity without QC and channel caveats"),
        ("Single neuron plus body orientation", "Neuron tracker or AFD_MTP", "Whole-worm centroid as neuron position"),
        ("Muscle structure", "Myocyte or nonstriated morphology mode", "One universal damage score"),
    ], [3100, 3260, 3000])
    before_you_film(doc)
    doc.add_heading("A standard analysis session", level=2)
    add_numbered(doc, [
        "Preserve raw data as read-only and analyze a copy or a designated results folder.",
        "Record biological metadata: strain, age, condition, temperature, assay geometry, and experimental unit.",
        "Record acquisition metadata: FPS, scale, exposure, channel roles, bit depth, and compression.",
        "Choose ROIs, intervals, orientation, and thresholds from visible biology, not to force an expected answer.",
        "Run automated detection and inspect QC overlays.",
        "Correct, relabel, or reject uncertain candidates; do not silently interpolate biological events.",
        "Export reviewed tables and retain provenance.",
        "Analyze at the correct inferential level: worm or plate, not frames, segments, or repeated events as independent N.",
    ])


def shared_principles(doc):
    doc.add_heading("Part II - Shared scientific and computational principles", level=1)
    doc.add_heading("Built from the biology back", level=2)
    doc.add_paragraph(
        "The ecosystem is not one generic computer-vision pipeline with different labels. Each tool begins "
        "with the biological feature to be captured, then chooses the minimal image evidence and human "
        "decisions needed to measure that feature. A pumping interval can be valid even when the rest of "
        "the movie is unusable. A plate-level orientation assay can be valid without maintaining worm "
        "identity. Basal slowing requires within-worm entry histories. These are biological distinctions, "
        "not interface preferences.")
    doc.add_heading("Why controls are exposed", level=2)
    add_bullets(doc, [
        "Frame rate and exposure determine which temporal events can be resolved.",
        "Spatial calibration determines whether distances and velocities are physical units or pixels.",
        "Object size, contrast, and tracking gates vary with magnification, illumination, strain, stage, and posture.",
        "ROIs encode assay geometry or suppress known artifacts; they are saved because they change the analyzed population.",
        "Before/after windows, buffers, and cadence limits encode a biological comparison and must remain auditable.",
        "Orientation and tissue-mode choices define anatomical coordinates and cannot always be inferred safely.",
        "Manual review exists where image evidence is ambiguous or biological identity matters."])
    doc.add_heading("Automation proposes; review establishes the analyzed set", level=2)
    doc.add_paragraph(
        "Experimental detectors should be understood as candidate generators. Orange, red, uncertain, "
        "collision, edge, short-coverage, saturation, and low-coherence states direct attention. They are "
        "not inconvenient records to delete. Reviewed outputs must remain linked to automated evidence and "
        "the original data.")
    doc.add_heading("Calibration and provenance", level=2)
    add_matrix(doc, ["Item", "Effect if wrong", "Required practice"], [
        ("FPS", "Frequency, rates, event durations, and velocities are wrong", "Use the acquisition setting or verified timestamps"),
        ("Micrometers per pixel", "Lengths, areas, and velocities are scaled incorrectly", "Use a stage micrometer or explicit two-point calibration"),
        ("Exposure", "Fast motion may be blurred and temporal precision overstated", "Record declared exposure and inspect blur"),
        ("Bit depth/compression", "Absolute fluorescence and fine morphology may be distorted", "Prefer raw data; preserve source-depth warnings"),
        ("Head-tail and dorsal-ventral orientation", "Signs, regions, and directional propagation can reverse", "Set orientation visibly and retain provenance"),
        ("ROI/interval", "The sampled population or time changes", "Save geometry and reviewed bounds"),
    ], [1900, 3560, 3900])
    doc.add_heading("QC is part of the measurement", level=2)
    add_bullets(doc, [
        "Segmentation quality: the mask represents the intended animal, tissue, cell, or egg.",
        "Identity quality: track continuity remains plausible across gaps and crossings.",
        "Temporal coverage: the recording includes enough valid time for the claimed frequency or event rate.",
        "Spatial coverage: the object is not clipped by the frame or excluded by an unrecorded ROI choice.",
        "Biological visibility: the relevant structure is genuinely discernible to a trained observer.",
        "Inferential independence: repeated frames, segments, encounters, and animals on one plate are not promoted to independent replicates."])


def locomotion(doc):
    doc.add_heading("Part III - Motor output: locomotion and sensory-guided behavior", level=1)
    tool(doc, "Track one worm (crawl, swim, burrow)", "Ready",
         "Create a signed, anatomically oriented time series of one visible worm's midline, posture, and movement.",
         "One animal can be followed at sufficient resolution and downstream segment-level kinematics are required.",
         [
             "Separate assay modes recognize that crawling, swimming, and burrowing differ in posture amplitude, speed, and expected resolution.",
             "The user marks the head because anterior-posterior sign cannot be inferred reliably from a symmetric mask.",
             "A user-drawn outline establishes animal size before automatic candidate selection; fragments and oversized persistent background structures outside the calibrated area envelope are rejected before they can become tracking hints.",
             "Camera motion is estimated from the coherent background before temporal-background comparison. Failed spine intervals have no fixed frame-count limit: reconstruction is attempted when trusted flanks differ by no more than half a body length in translation and satisfy the posture-agreement gate.",
             "Signed curvature is preserved rather than reduced to unsigned bend magnitude."
         ],
         [("Assay mode", "Changes tracking and resolution expectations for crawl, swim, or burrow."),
          ("FPS, scale, exposure", "Define temporal resolution, physical units, and blur limits. Scale must be a positive calibrated value; zero or unknown is refused."),
          ("Head click and outline", "Anchor anatomical orientation, body size, and segmentation."),
          ("Manual frame fixes", "Repair tracking failures without discarding the whole recording.")],
         [
             "The tracker segments the chosen worm, orders a head-to-tail midline, and carries the mask and midline forward through the movie.",
             "The initial trusted frame seeds tracking in both temporal directions. Conserved length, area, whole-spine continuity, and registered residual motion constrain identity; isolated endpoint jumps may be stabilized.",
             "Any two-sided failed interval may be reconstructed between trusted flanking spines when normalized translation and posture are bridgeable. With N internal anchors, the interval becomes N+1 independently reconstructed subintervals. Previously inferred spines are never reused as anchors after a new manual correction. This all-anchor rule is also used by the anterior-neuron/body reviewer, RGBCaMP Fiji manual midlines, and pBoc outline landmarks. Raw geometry is retained, and repaired frames export inferred_between_neighbors provenance.",
             "For a bounded correction, place the viewer on a trusted spine and press b, then move to a trusted ending spine and press e. The boundary spines remain anchors and only frames strictly between them are reconstructed. Bounded-edit mode remains active: press f on any intervening frame to draw additional anchors. Each new f anchor invalidates the prior interpolation, preserves every manual and boundary anchor, and rebuilds only the subintervals between them. Press a to visit a suggested midpoint and c when finished. No f correction in this mode can retrack outside b...e.",
             "Press w to save progress while reviewing, q or close to save and exit without finalizing, and s to finalize the CSV. Reopening any frame from the same numbered recording offers to restore the session. Selecting one numbered still loads only its shared recording prefix and numeric suffix."
         ],
         ["Tracking CSV with time, centroid, midline, signed curvature, body geometry, and provenance.",
          "Use the CSV as the input to Kinematics analysis or Single-worm swimming analysis."],
         ["Coils, self-contact, overlaps, clipping, low contrast, abrupt focus changes, and motion blur can break the midline.",
          "An incorrect head assignment reverses anatomical position and directional signs.",
          "A zero or blank micrometres-per-pixel value is not an unknown-scale mode; it is an invalid calibration and analysis is stopped before loading.",
          "Thousands of separate images on a network share remain limited by server and file-open latency. Copy one selected sequence to a local SSD for tracking when practical, then retain the reviewed output with the original data.",
          "Global brightness flicker or auto-gain can destabilize segmentation. Camera registration corrects translation, not exposure changes, rotation, severe blur, or non-rigid field distortion. Lock exposure and gain during acquisition whenever possible.",
          "A mode cannot recover spatial or temporal information absent from the recording."],
         ["Correct only visibly wrong frames; do not redraw toward an expected waveform.",
          "If the animal repeatedly leaves the field, shorten to a defensible interval.",
          "Increase acquisition rate for high-frequency swimming; improve contrast before relaxing segmentation excessively."])

    doc.add_heading("Supervised segmentation workbench", level=3)
    doc.add_paragraph(
        "The workbench defines the binary object mask; it does not replace anatomical orientation, "
        "component identity, midline extraction, conserved-length/area QC, manual correction, or "
        "temporal reconstruction. Press g in the DIC tracker review window to open it. The workbench "
        "runs in an isolated window; after Accept + Lock and close, the tracker reloads the recorded "
        "settings and retracks automatically.")
    add_callout(
        doc, "Default-preservation rule",
        "An accepted map is opt-in. With no accepted map, the original DIC detector is unchanged. "
        "RGBCaMP retains its Fiji/manual-midline default and cannot receive this DIC map through the "
        "photometry firewall.")
    doc.add_heading("Choose frames and ranges", level=3)
    add_numbered(doc, [
        "Place the Good-frame preview on a representative successful frame and the Bad-frame preview on a difficult frame. Both previews use the same current settings.",
        "Use Previous exact, Next exact, the left/right arrow keys, or Jump to frame to load an exact source frame. The coarse sliders remain useful for rapid scanning.",
        "If one recipe works throughout, leave the range table empty: the global controls apply to every frame.",
        "For a bounded change, mark the inclusive start and end source frames, adjust the controls, and choose Apply settings to range. Frames outside explicit ranges use the global controls.",
        "Inspect both sides of every range boundary. Overlapping ranges are refused; a newly applied intersecting recipe replaces the older intersecting recipe.",
        "Choose Accept + Lock, confirm the review reminder, and close the workbench. Do not finalize the tracker CSV until the new flagged-frame count and representative midlines have been reviewed."
    ])
    doc.add_heading("Escalation ladder: least to more intervention", level=3)
    add_matrix(doc, ["Strategy", "When and how to use it"], [
        ("1. Global gray threshold", "Start here. In dark mode adjust High; in bright mode adjust Low; band mode uses both. This is cheapest and easiest to interpret."),
        ("2. Morphological cleanup", "Increase Close gradually to join small gaps. Fill body repairs interior holes. Minimum object area discards small components. Excessive cleanup can join the worm to bacterial structures."),
        ("3. Local contrast", "Choose the local_contrast feature when absolute illumination varies but DIC body texture/edges remain stronger than the surrounding agar. Bacterial texture can also respond."),
        ("4. Background difference", "Choose difference to compare the frame with a registered temporal-median background. It is useful for a translating worm on a stationary lawn but weakens during prolonged stillness."),
        ("5. Local-adaptive mode", "Use when illumination varies broadly across the field. It is more aggressive and may select local bacterial texture."),
        ("6. Space-time ROI", "Draw a polygon around a localized illumination problem, then assign its inclusive frame interval and local threshold. Ignore a bad region if the target never enters it."),
        ("7. Identity safeguards", "Previous-mask overlap prevents a jump to a spatially separate bacterium or newcomer. Camera registration moves the expected centroid/mask with coherent camera translation. A genuine cut clears inherited identity."),
    ], [2600, 6760])
    doc.add_heading("Interpret the preview and QC correctly", level=3)
    add_bullets(doc, [
        "The red overlay is the proposed foreground mask, not a guarantee that the final midline, head, tail, length, or area is correct.",
        "Low is intentionally disabled in dark mode and High is intentionally disabled in bright mode; both are active in band mode.",
        "Length and area conservation are evaluated by the tracker against the user's initial head click and outline. They are downstream QC, not threshold-preview controls.",
        "A lower flagged-frame count demonstrates internal consistency, not anatomical truth. Spot-check head/tail polarity and midlines across early, middle, late, difficult, and range-boundary frames.",
        "If no plausible overlapping component exists, the tracker should flag the frame rather than switch identity. A terminal frame without a trusted right neighbor cannot be safely reconstructed from both sides; correct it manually only when the worm is visibly resolvable.",
        "Record all non-default ranges, thresholds, features, morphology, ROIs, overlap, registration, manual anchors, and unresolved frames in methods/provenance."
    ])

    tool(doc, "Kinematics extractor (Fiji)", "Ready; Fiji companion extractor",
         "Provide a manually supervised Fiji route to an approved midline and kinematics CSV.",
         "The scientist prefers Fiji/ImageJ interaction or needs direct midline approval in a stack.",
         ["Manual midline approval is retained because anatomical errors propagate into every segment metric.",
          "Extraction is separated from downstream biological summarization so raw measurements remain reusable."],
         [("Midline/outline approval", "Defines the anatomical object used for segment measurements."),
          ("Stack and calibration", "Determine time and physical scale."),
          ("Frame corrections", "Allow biologically informed rescue of isolated failures.")],
         ["The Fiji tool follows the worm through the stack, presents the midline for approval, and exports a structured CSV compatible with the Python analysis layer."],
         ["Approved kinematics CSV and Fiji-side visual evidence."],
         ["Fiji version/plugin installation differences; incorrect calibration; ambiguous head-tail; coils and self-contact.",
          "Manual extraction can introduce scorer variability; document the scorer and review standard."],
         ["Use the same approval rubric across conditions and keep blinded scoring where possible."])

    tool(doc, "Kinematics analysis", "Ready",
         "Convert an extracted single-worm time series into biologically interpretable body-wave, locomotion, foraging, and dampening measurements.",
         "A validated extractor CSV exists and the question concerns waveform, propagation, segment behavior, or locomotor output.",
         ["Analysis is downstream of extraction so measurement definitions can change without retracking.",
          "Segment-level curvature is retained because posterior dampening and head foraging cannot be inferred from centroid speed.",
          "QC-qualified intervals are summarized rather than treating all frames as equally valid."],
         [("Input CSV", "Selects the recording and its embedded acquisition/QC metadata."),
          ("Analysis options exposed by launcher", "Allow the scientist to choose outputs appropriate to the assay rather than computing unsupported claims.")],
         ["The pipeline reads signed curvature and movement, constructs curvature kymographs, identifies body-wave descriptors, and produces locomotion, foraging, and posterior-dampening summaries."],
         ["Summary tables, segment/timecourse tables, kymographs, head-swing traces, and dampening figures."],
         ["Poor extractor output; wrong head-tail orientation; missing FPS/scale; too little usable time; self-contact that shortcuts curvature.",
          "Frequency estimates from short or discontinuous traces may be unstable."],
         ["Inspect the browser before interpreting a summary; restrict analysis to valid continuous runs when necessary."])

    tool(doc, "Kinematics browser", "Ready; results viewer",
         "Let the scientist inspect one recording's derived kinematics rather than relying only on summary numbers.",
         "A kinematics analysis folder exists and segment/time relationships need visual checking.",
         ["A browser is a scientific control: segment selection and kymograph inspection expose artifacts hidden by means."],
         [("Selected segments", "Focuses curvature traces on anatomically relevant regions."),
          ("Recording/result set", "Preserves one-to-one linkage with source analysis.")],
         ["The browser loads result tables and figures, displays signed-curvature kymographs, and allows selected segment traces to be explored."],
         ["Interactive views; no new biological measurements unless explicitly exported."],
         ["Missing/incomplete result files; confusing recordings; overinterpretation of visually selected segments."],
         ["Choose segments before inspecting condition labels when exploratory selection could bias conclusions."])

    tool(doc, "Single-worm swimming analysis", "Ready",
         "Measure swimming frequency, occupancy, amplitude, phase organization, and usable coverage from one high-resolution swimming track.",
         "A Track one worm CSV was acquired in swimming mode and segment-resolved swimming metrics are needed.",
         ["Swimming is analyzed from curvature, not centroid oscillation alone.",
          "Usable coverage and contiguous runs are reported because an average frequency can conceal missing or fragmented data.",
          "Segment amplitude and phase retain where and how the wave changes along the body."],
         [("Input swimming CSV", "Ensures the correct assay mode and metadata."),
          ("Run/quality thresholds", "Define how much continuous evidence is required for a frequency estimate.")],
         ["The tool identifies valid contiguous swimming runs, estimates oscillation frequency, computes segment amplitude profiles and phase relationships, and records frame-level state/QC."],
         ["Swimming summary; contiguous-run frequencies; amplitude profile; frame QC/state."],
         ["Short bouts; low FPS; blur; poor midlines; coils; mixed crawling/burrowing in the selected recording.",
          "A whole-recording mean may be biologically misleading when modality changes."],
         ["Use the population modality tool for mixed-population movies; split or relabel mixed single-worm bouts before combining them."])

    tool(doc, "Population tracking", "Experimental",
         "Track many animals and quantify reviewed swimming, crawling, burrowing, and uncertain bouts independently within one movie.",
         "Multiple worms are visible at low magnification and population-level modality occupancy or kinematics are desired.",
         ["Identity linking uses predicted motion and flags close competing assignments.",
          "The same lazy source reader accepts MP4/AVI/MOV/MKV/WebM video, multipage TIFF stacks, and folders of common sequential images.",
          "Compressed video is decoded sequentially: one streaming pass samples the background and one streaming pass performs analysis, avoiding whole-movie loading and repeated from-start decoding.",
          "A 25-point oriented spine replaces the earlier ellipse-axis frequency proxy.",
          "Swimming, crawling, and burrowing use combined frequency, C/S/W curvature topology, speed, wave lag, and temporal persistence; frequency alone is insufficient.",
          "Four-second overlapping proposals are smoothed into bouts, but every bout remains a human-review candidate.",
          "ROIs are optional: full-frame analysis is the default; include/exclude geometry exists for assay boundaries and static artifacts."
         ],
         [("FPS and scale", "Set frequency and physical velocity."),
          ("Minimum/maximum object area", "Adapt segmentation to magnification, stage, and posture."),
          ("ROI mode: none/include/exclude", "Keep full-frame default or restrict/suppress detections in oval, rectangle, or polygon regions."),
          ("Measure a worm / Mark all animals", "Click one animal, or every animal, to set the area gates and the link distance from what the detector actually measures rather than from a guess. Marking every animal also gives the review an expected count to check against."),
          ("Max link (px/frame)", "How far one animal may travel between frames, in source pixels. Too large and the tracker welds separate animals into one track."),
          ("Spine skeleton", "Standard (historical default) or connected thinning. The two are not interchangeable: spine, curvature and bend frequency depend on the choice, which is recorded in analysis_metadata.json."),
          ("Manual track acceptance", "Controls which identities enter bout review."),
          ("Track editing", "Stitch any number of fragments, split at a frame, delete, lasso-select, add missing points, and fill gaps. Placed points carry identity only."),
          ("Bout confirm/relabel/reject", "Establishes the final biological modality labels.")],
         ["Choose Movie / stack for a video or multipage TIFF, or Image folder for sequential frames. A movie already loaded in the Hub is handed directly to this tool.",
          "A median background isolates moving objects; connected components are linked by constant-velocity prediction. Each mask is skeletonized and resampled to an oriented spine.",
          "Overlapping windows combine signed midbody curvature frequency, C/S/W scores, centroid speed, posterior wave lag, collision evidence, and spine coverage. A visual player shows the original frames, spine, centroid, and trajectory before final assignment.",
          "Review happens inside the module window: tracks are drawn over the movie frames and stay synchronised as you scrub or play, each accepted animal in its own colour, so a track can be checked against the animal it claims to follow.",
          "Positions you place by hand are flagged manual_point and are excluded from speed, coverage, frequency and curvature. They carry identity across frames the detector missed and are never measured; track summaries report detected_frames and manual_points separately."],
         ["Frame-level tracks/spines; track summaries; modality-window evidence; pending and reviewed bouts; per-modality time, bout count, speed, and frequency; ROI and analysis provenance."],
         ["Merged worms, crossings, stationary animals absorbed into the background, debris that moves, frame edges, poor skeletons, and low temporal resolution.",
          "Burrowing and crawling frequency ranges overlap; posture and wave evidence can still be ambiguous.",
          "Centroid-based ROI inclusion can omit a worm whose body crosses a boundary while its centroid remains outside."],
         ["Exclude persistent debris with an ROI; include only the assay chamber when the field contains irrelevant regions.",
          "Adjust area gates only after inspecting masks; too narrow loses posture extremes, too broad admits merges.",
          "Retain uncertain bouts or report their fraction; do not force labels to make percentages sum neatly."])

    tool(doc, "wrMTrck (Fiji)", "Ready as an external Fiji plugin entry",
         "Provide a familiar many-worm tracking option inside Fiji.",
         "The experiment matches wrMTrck assumptions and its outputs are sufficient.",
         ["The Hub exposes the established plugin without pretending it shares the Lab Tools QC/provenance model."],
         [("Plugin thresholds/settings", "Adapt detection and linking to the recording; definitions depend on the installed wrMTrck version.")],
         ["Run wrMTrck from Fiji using its documented workflow and preserve its settings with results."],
         ["Plugin-defined trajectories and summaries."],
         ["Version sensitivity; limited integration with Lab Tools metadata; merges/crossings; threshold dependence."],
         ["Record Fiji and plugin versions. Do not compare wrMTrck and Lab Tools outputs as if definitions were identical."])

    tool(doc, "Population basal slowing", "Experimental",
         "Measure paired within-worm changes in velocity and bending before, during, and after OP50 lawn encounters, including repeated encounters.",
         "A population disperses from a release droplet toward visible food lawns and individual histories can be reconstructed.",
         ["The start ROI is a release gate: clustered animals are ignored until a first observed or inferred exit.",
          "Tracks remain active after later re-entry and through repeated lawn encounters.",
          "Entry uses a configurable fraction of worm area inside the lawn, not centroid crossing alone.",
          "Pre-entry frames must satisfy an outside buffer to avoid mixing boundary contact into baseline.",
          "Conservative tracklet stitching uses position and aligned direction and preserves every inferred join.",
          "Missing curvature frequency does not invalidate an otherwise usable paired speed comparison."
         ],
         [("Start and lawn ROIs", "Encode the actual release geometry and each food lawn."),
          ("Before/after windows", "Balance biological timescale against available unambiguous frames."),
          ("Outside buffer", "Defines a clean pre-contact baseline."),
          ("Fraction inside for entry", "Matches how much body contact constitutes biological entry."),
          ("Area/link/stitch gates", "Adapt detection and identity continuity while limiting false merges."),
          ("Track and event review", "Determines which identities and encounters are scientifically usable.")],
         ["Bright worms are detected against a low temporal background percentile. Masks yield centroids and 25-point spines. Tracks are linked and conservatively stitched, then evaluated against release and lawn geometry.",
          "Every entry receives paired windows, encounter order, elapsed times, cumulative prior exposure, residence time, post-exit measurements, and explicit QC reasons."],
         ["Reviewed tracks and stitches; frame-level position/spine/ROI state; paired entry events; reviewed events; descriptive summary; exact ROI and acquisition metadata."],
         ["Identity ambiguity at crossings; segmentation changes on food; boundary and edge effects; insufficient pre-entry space/time; collision; repeated encounters from the same worm.",
          "Events are repeated observations, not independent animals. Absolute speed depends on calibration."],
         ["Shorten exploratory windows only with a documented rationale; preserve the original analysis.",
          "Use mixed-effects/repeated-measures models for encounter histories.",
          "Visually inspect apparent accelerations for centroid shifts caused by segmentation on OP50."])

    tool(doc, "Population orientation (plate state)", "Ready for identity-free population-state configuration",
         "Measure orientation toward/away from a stimulus at the plate level without requiring reliable individual identities.",
         "The biological question concerns occupancy, angular/radial distribution, arrival, or a plate-level directional resultant.",
         ["Identity-free pixels/blobs are preferred when crowding or crossings make individual paths unreliable.",
          "The plate, not each worm, is the inferential unit.",
          "Stimulus, control, release positions, and scale are explicitly defined because assay geometry determines every angle and distance.",
          "A per-worm path configuration remains gated until acquisition supports it."],
         [("Plate ID", "Preserves the independent experimental unit."),
          ("FPS and two-point scale", "Set temporal and spatial axes."),
          ("Stimulus/control/release positions", "Define biological coordinate geometry."),
          ("ROIs and detection settings", "Match the arena and visible population state.")],
         ["Median background subtraction estimates occupancy and descriptive blobs. Radial and angular pixel distributions, arrivals, timecourses, and the plate resultant are computed without claiming persistent identity."],
         ["ROI occupancy, radial/angular distributions, plate timecourse, and one plate-level resultant."],
         ["Uneven illumination, stationary animals disappearing into background, debris, crowded masks, inaccurate stimulus coordinates, and pseudo-replication.",
          "Configuration 2 path analysis is not validated for the current framing."],
         ["Use multiple independent plates; do not pool worms across plates as N.",
          "Re-frame and reacquire if the question truly requires individual paths."])

    tool(doc, "Combine orientation plates", "Ready",
         "Perform circular statistics across independent plate resultants.",
         "Two or more plate_resultant.csv files represent independent experimental units.",
         ["The tool combines plates rather than worms to prevent pseudo-replication."],
         [("Selected plate-resultant files", "Define the independent replicates and groups entering the comparison.")],
         ["The aggregator reads one resultant per plate and applies circular summaries/comparisons at the plate level."],
         ["Across-plate circular statistics and combined summaries."],
         ["Too few plates; mixing non-independent plates; inconsistent stimulus geometry or acquisition; pooling incompatible conditions."],
         ["Predefine exclusion rules at the plate level and retain each plate's original analysis folder."])


def rhythmic(doc):
    doc.add_heading("Part IV - Motor output: rhythmic programs", level=1)
    tool(doc, "Pharyngeal pumping", "Ready; new detector requiring manual-count validation",
         "Quantify reviewed pump events only during intervals in which the pharynx is visibly scoreable.",
         "The pharynx/grinder is resolved and at least one relatively still, in-focus interval exists.",
         ["The user chooses a usable interval rather than allowing invisible frames to dilute the rate.",
          "A close pharyngeal ROI localizes the motion source.",
          "Tracking/focus/motion gates reject unusable frames, and the detection threshold remains reviewable."],
         [("FPS", "Converts events and intervals to rate and time."),
          ("START/END", "Defines the biologically visible scoring interval."),
          ("Oval pharyngeal ROI", "Restricts motion measurement to the intended structure."),
          ("Detection-threshold slider", "Lets the reviewer match peaks to visually resolved pumps.")],
         ["The ROI is tracked through the approved interval. Excessive motion, weak tracking, poor focus, and insufficient contrast are excluded; candidate motion peaks are displayed for threshold review."],
         ["Approved interval; reviewed event times/signals; pump count; pumping rate; median interpump interval."],
         ["Head movement mimicking pumps; grinder out of focus; obstruction; low FPS; ROI drift; extending through unscoreable periods.",
          "Automatic counts are not yet a substitute for manual validation."],
         ["Prefer 25-40 FPS and at least 30 seconds; analyze multiple visible intervals separately rather than bridging invisibility."])

    tool(doc, "Defecation cycle analysis", "Experimental alpha",
         "Propose and review posterior body contraction (pBoc) events and cycle cadence.",
         "One complete pBoc can be identified with baseline, peak, and recovered outlines; the posterior is visible around later events; and every additional worm can be annotated over its visible interval.",
         ["Motion is decomposed into axial versus dorsoventral components because pBoc is posterior axial motion toward the head followed by recovery.",
          "Equivalent anterior motion argues against a specific posterior contraction.",
          "Three outlines calibrate the last full-length frame before contraction, the minimum-length peak, and the first fully recovered frame. The observed shortening is expected to be about 5-8%, but recording-specific measurements define the model.",
          "The calibration navigator displays the movie directly; users scroll to and outline the three landmarks without knowing frame numbers beforehand. Calibration is saved and resumable.",
          "All three adult outlines define admissible target length and area. Substantially smaller larvae are rejected as identities rather than retained as shortened pBoc geometry.",
          "Length may vary continuously between baseline and peak while area should remain approximately conserved. Raw measurement geometry is kept separate from regularized tracking geometry.",
          "Within the registered worm mask, the detector measures the fraction of textured posterior and anterior pixels whose residual motion is primarily axial rather than transverse.",
          "Cycle-period limits order reviewer attention but never accept, reject, or establish a biological period.",
          "Rhythm statistics are withheld until at least ten manually accepted cycles are followed continuously.",
          "Additional worms are moving identities, not static exclusion ROIs. Seed each with a segmented centerline at its first clear frame and declare its last visible frame.",
          "If a distractor cannot be tracked or approaches the target, affected target frames are unusable. The software does not guess identity through overlap."],
         [("FPS, scale, exposure", "Define temporal/physical resolution and blur."),
          ("Minimum/maximum plausible cycle period", "Flag close candidates and long gaps for review; defaults 30-90 seconds."),
          ("Three calibration frames and outlines", "Trace head, tail, and the complete worm at pre-pBoc baseline, minimum-length peak, and first full recovery."),
          ("Moving-distractor episodes", "For each non-target worm, define its first clear frame, segmented centerline, and last visible frame; use a separate episode after re-entry."),
          ("Manual event decisions/cycle limits", "Establish the accepted event series.")],
         ["In the calibration navigator, scroll to the last full-length frame before one clear pBoc, its minimum-length peak, and the first frame after full length returns. On each selected frame click head, tail, and trace the entire outline.",
          "The tool measures baseline/peak/recovered length and area, shortening fraction, contraction and recovery duration, pixel and physical rates, and posterior-versus-anterior axial-pixel participation. Implausible shortening, recovery disagreement, or area loss produces calibration warnings.",
          "Before analysis, scrub the entire recording. At each distractor's first clear frame, choose Add moving distractor, draw a segmented line end-to-end, set the last visible frame, and save. If no other worms occur, explicitly save No distractors.",
          "The engine tracks each annotated distractor independently, excludes its moving mask from target segmentation, and then analyzes posterior and anterior target motion.",
          "The accepted calibration event seeds bidirectional target tracking. Later candidates combine axial flow, calibrated shortening, area conservation, axial-pixel participation, tracking confidence, and cadence priority.",
          "In visual review, distractors are magenta. Pink/red distractor overlays and target identity warnings mark frames that cannot support measurement."],
         ["Candidate/review tables; raw and combined motion scores; calibrated contraction fraction; area error; posterior/anterior axial-pixel fractions; camera shifts; geometry provenance. Period, IDI, and IDI CV remain unavailable until validation criteria are met."],
         ["Turns, whole-body translation, focus changes, incomplete recovery, tail occlusion, short recordings, cadence-driven confirmation bias, an unannotated extra worm, a distractor lost in the lawn, and prolonged worm-worm overlap."],
         ["Treat cadence as a search aid only. Inspect the underlying motion around every proposed and missing event.",
          "Annotate every non-target worm, including one present only briefly. A re-entering worm needs a new episode. Do not shorten an episode to hide failed tracking.",
          "When overlap prevents identity assignment, reacquire with better animal separation; threshold tuning cannot make identity observable."])

    tool(doc, "Endpoint egg counting", "Experimental",
         "Count resolved eggs in a static image with scale-aware candidate detection and mandatory correction.",
         "Eggs are spatially resolved and the image can be calibrated.",
         ["Expected egg dimensions are physical (approximately 50 x 30 um), so two-point calibration precedes detection.",
          "Automatic candidates are never final; users can toggle false positives and add missed eggs."],
         [("Two-point scale calibration", "Maps biological egg size to pixels."),
          ("Candidate settings/visual corrections", "Adapt to contrast, overlap, genotype, developmental state, and debris.")],
         ["The detector identifies objects compatible with calibrated egg size/shape, overlays candidates, and saves only after visual review."],
         ["Reviewed egg list/count, candidate evidence, calibration, and summary."],
         ["Overlapping eggs; embryos with atypical shape; debris of similar size; blur; uneven illumination; incorrect calibration; eggs at edges."],
         ["Use multiple focal planes or improved imaging when eggs overlap in depth; do not solve unresolved overlap by threshold tuning alone."])

    tool(doc, "Dynamic egg laying", "Experimental",
         "Detect and review when persistent new eggs appear during a recording.",
         "Eggs remain visible after deposition and acquisition resolves their appearance over time.",
         ["Persistence distinguishes a newly laid egg from transient noise or a moving reflection.",
          "Timing candidates require review because worm motion or occlusion can delay first visibility."],
         [("FPS", "Sets timing precision."),
          ("Two-point scale", "Constrains egg size."),
          ("Persistence and detection settings", "Balance transient rejection against missed/occluded eggs."),
          ("Manual candidate timing review", "Defines the accepted laying event.")],
         ["Frame-to-frame candidates are linked as persistent new objects. Review overlays expose first detection and persistence before saving event times."],
         ["Frame QC, egg-object tracks, automatic candidates, reviewed event timing, and metadata."],
         ["Egg occluded at deposition; worm overlaps; drifting debris; illumination jumps; moving eggs; insufficient post-event duration; incorrect scale."],
         ["Interpret first visible frame as an observation bound when deposition itself is occluded."])


def physiology(doc):
    doc.add_heading("Part V - Physiology: calcium and cellular activity", level=1)
    tool(doc, "RGBCaMP extractor (Fiji)", "Ready",
         "Extract spatially registered blue, green, and red signals with body curvature and kinematics from a freely moving worm.",
         "A multichannel recording contains the intended indicators and the worm can be tracked and segmented.",
         ["The moving body is divided into 24 hemisegments per side so fluorescence is interpreted in anatomical coordinates.",
          "Eigenworm-constrained midlines reduce implausible geometry.",
          "Extraction stops at raw, QC-rich data; normalization and inference remain downstream.",
          "Channel roles are explicit: green cytoplasmic GCaMP, red mito RCaMP/pharynx mCherry, blue ER indicator."],
         [("Channel assignment", "A wrong mapping changes the biological meaning of every signal."),
          ("Tracking/midline parameters", "Adapt to visibility while preserving plausible anatomy."),
          ("Calibration and acquisition metadata", "Determine spatial units and saturation/bit-depth interpretation.")],
         ["The Fiji plugin tracks the worm in DIC, fits the midline, divides the body into hemisegments, and extracts channel intensities, ratios, derivatives, curvature, and movement with QC fields."],
         ["Per-frame, per-segment CSV used by RGBCaMP analysis."],
         ["8-bit compression; saturation; channel bleed-through; pharyngeal/head contamination; missing scale; coils; partial body; signal outside the mask.",
          "Absolute intensity from compressed 8-bit sources is not fully quantitative."],
         ["Prefer raw high-bit-depth recordings. Preserve filename metadata grammar and inspect saturation/QC before downstream analysis."])

    tool(doc, "RGBCaMP analysis", "Ready",
         "Transform extraction CSVs into QC-aware relative fluorescence, spatial summaries, kinetics, coupling, and condition-level outputs.",
         "One or more RGBCaMP extraction CSVs exist and the question concerns calcium dynamics or calcium-behavior coupling.",
         ["dF/F0 and ratios are emphasized for 8-bit sources; absolute resting intensities remain caveated.",
          "Frames with coil/area/size/short-length/partial/low-evidence or outside-body leakage are rejected.",
          "Self-approach is a curvature caveat rather than a whole-frame photometry rejection.",
          "Red and green are masked in segments 0-7 where pharyngeal mCherry/head GFP contaminate myocytes 1-8; blue remains available.",
          "Wave propagation uses phase gradients because adjacent delays are below one frame at low sampling rates.",
          "The animal, not segment or transient, is the inferential unit."],
         [("Input CSV/folder", "Defines one recording or batch."),
          ("Filename metadata", "Encodes genotype, age, RNAi target, control, and flagged quality."),
          ("QC policy", "Controls frame retention; changes must be reported."),
          ("Channel/region/metric selection", "Matches the biological hypothesis and known contamination."),
          ("Grouping/statistical choices", "Must respect animal-level replication.")],
         ["The pipeline loads and QC-filters data, computes dF/F0 and ratios, summarizes reliable body regions, detects transients, estimates release/reuptake, coupling, intersignal timing, phase-gradient waves, and cycle averages.",
          "Condition statistics operate on per-worm summaries using nonparametric effect sizes and FDR; mixed models are the appropriate extension for nested datasets."],
         ["QC report; per-worm/segment summaries; kymographs; kinetics; region split; contraction state; coupling; wave propagation; cycle averages; condition statistics."],
         ["Saturation, 8-bit sources, incorrect channel mapping, filename metadata not parsing, missing scale, low coherence, decay fits at bounds, too few animals, and pseudoreplication.",
          "Head-region green/red metrics are unreliable by design and masked."],
         ["Report phase-gradient speed only with meaningful coherence (approximately >=0.5).",
          "Treat decay constants pinned at the 20-second bound as rejected.",
          "Rename future files consistently: genotype_day_RNAi_condition.csv."])

    tool(doc, "RGBCaMP browser", "Ready; results viewer",
         "Inspect recording-level calcium, kinematics, QC, and derived figures before comparison.",
         "An RGBCaMP analysis folder exists.",
         ["Interactive inspection prevents summary statistics from hiding saturation, missing regions, rejected frames, or desynchronized signals."],
         [("Recording/result selection", "Links views to the intended animal."),
          ("Channel/metric/segment view", "Lets the scientist inspect the relevant biological signal.")],
         ["The browser reads output tables and figures and organizes recording QC, timecourses, spatial patterns, and derived metrics."],
         ["Interactive review; no independent replication is created by exploring many segments."],
         ["Incomplete result folders; visual cherry-picking; ignoring QC warnings."],
         ["Define primary channels/regions before condition comparison."])

    tool(doc, "Neuron tracker", "Ready",
         "Track an anterior sensory neuron soma and the worm's body orientation through a movie.",
         "The soma is visible and its position relative to the moving animal is the biological measurement.",
         ["The user marks several points inside the soma and outlines the worm because cell and body identity require anatomical supervision.",
          "Neuron position is interpreted with body orientation, not as an image-space centroid alone.",
          "Flagged frames can be manually repaired.",
          "Positive scale and exposure are validated before the recording is loaded. Separate image files use bounded parallel decoding into one preallocated array."],
         [("Soma clicks", "Initialize cell size/location."),
          ("Worm outline", "Establish body reference and orientation."),
          ("Tracking/QC corrections", "Repair loss, ambiguity, or outline failure."),
          ("FPS, scale, exposure", "Set temporal and physical interpretation.")],
         ["The tracker follows the soma and worm mask, records provenance for automatic and corrected frames, and exports neuron/body geometry through time."],
         ["Neuron trajectory, body orientation, QC/provenance, and review views."],
         ["Cell disappears, neighboring fluorescent objects, photobleaching, body overlap, focus loss, ambiguous anterior anatomy, and outline clipping."],
         ["Use manual correction for isolated errors; reacquire when the soma is unresolved for extended periods.",
          "For very large image sequences on a shared drive, copy the recording to a local SSD for analysis when practical; keep results linked to the archived original."])

    tool(doc, "AFD_MTP (Fiji)", "Ready; Fiji companion",
         "Provide a Fiji implementation of anterior sensory-neuron/body tracking.",
         "The scientist needs the established Fiji workflow or comparison with prior AFD_MTP analyses.",
         ["Maintaining a Fiji route supports legacy data and direct stack interaction; outputs should not be assumed identical to the Python tracker."],
         [("Fiji tracking parameters and manual corrections", "Adapt cell/body tracking to the recording and plugin version.")],
         ["The plugin tracks the fluorescent neuron and body reference in Fiji and exports its measurements."],
         ["Plugin-specific neuron/body tracking outputs."],
         ["Fiji/plugin version differences; gap handling; cell loss; competing objects; calibration errors."],
         ["Record the plugin version and avoid mixing implementations within an experiment without a parity study."])


def morphology(doc):
    doc.add_heading("Part VI - Anatomy and morphology", level=1)
    tool(doc, "Myocyte morphometry (Fiji)", "Ready",
         "Measure striated myocyte shape and actin-fiber waviness.",
         "Fluorescence resolves myocyte boundaries/fibers and the question concerns structural degeneration.",
         ["Shape and fiber waviness are retained as interpretable features rather than collapsed immediately into one score."],
         [("ROI/segmentation thresholds", "Adapt to fluorescence, magnification, and tissue quality."),
          ("Scale", "Converts morphology to physical units.")],
         ["The Fiji macro segments selected muscle structures and computes geometric and waviness measurements."],
         ["Per-object morphometry and visual segmentation evidence."],
         ["Uneven staining; saturation; out-of-focus fibers; overlapping structures; inconsistent ROI placement; threshold dependence."],
         ["Blind scorers to condition and use consistent anatomical sampling across animals."])

    tool(doc, "Nonstriated muscle degeneration", "Experimental",
         "Quantify tissue-specific structural features in pharyngeal, uterine, somatointestinal, or anal-depressor muscle.",
         "A raw fluorescence image, calibration, tissue ROI, and anatomical orientation are available.",
         ["Separate tissue modes recognize different architectures and force directions.",
          "The anal-depressor mode requires proximal attachment and distal insertion clicks and reports a force-vector angle in worm coordinates.",
          "Orientation defaults may persist for convenience, but per-image provenance is exported.",
          "The composite damage score is deliberately blank until WT, dystrophic, and rescue reference sets calibrate it."],
         [("Tissue mode", "Selects biologically appropriate features and geometry."),
          ("Scale calibration", "Converts dimensions to physical units."),
          ("Tissue ROI/segmentation controls", "Define the anatomical structure and accommodate staining variation."),
          ("Body orientation", "Defines worm coordinates and directional signs."),
          ("Attachment/insertion clicks", "Define anal-depressor force-vector geometry.")],
         ["The tool segments the chosen tissue inside a user-defined ROI, computes interpretable structural/geometric features, and saves an overlay plus acquisition/orientation provenance."],
         ["Segmentation overlay; tissue-specific morphology features; force-vector geometry where applicable; no uncalibrated composite damage score."],
         ["Wrong tissue mode; orientation reversal; saturated or uneven signal; incomplete ROI; ambiguous attachments; segmentation of background or adjacent tissue.",
          "A future composite score could overfit if calibrated on too few reference animals."],
         ["Analyze raw component features until reference calibration is sufficiently broad; keep orientation evidence with every image."])


def mechanosensation(doc):
    doc.add_heading("Part VII - Mechanosensation and evoked response", level=1)
    tool(doc, "Evoked mechanosensation and habituation", "Experimental",
         "Quantify how a single worm responds to a mechanical stimulus (tap, nose touch, gentle or harsh body touch) and how that response habituates across repeated stimuli, preserving trial order, denominator, and each animal's own baseline.",
         "You have a recording of a stimulated worm and want reversal onset latency, reversal velocity, duration and distance, head-bend amplitude, and a stop-versus-reverse classification, optionally across a habituation series.",
         ["Tracking and scoring are one flow: the tool tracks the movie and auto-loads the kinematics, so there is no manual CSV hand-off.",
          "Stimulus timing is a time marker, typed or dropped on the movie; the pick artifact is handled by skipping non-finite frames while latency stays anchored to the entered stimulus time (an optional blackout window excludes the pick-in-view interval).",
          "A response is not always a reversal: an animal can stop or pause without reversing, which is tracked as its own outcome.",
          "Harsh-touch location matters: anterior stimuli evoke a reversal, posterior stimuli a forward escape (accelerated forward crawling); both are scored.",
          "A design selector separates a single trial, a habituation series on the same animal (dependent trials, order and inter-stimulus interval preserved), and independent animals (independent replicates); the plate is the unit of replication.",
          "A spontaneous mode detects every reversal with no stimulus at all."],
         [("Stimulus type", "Nose, gentle, or harsh touch, or population tap."),
          ("Stimulus location", "Anterior vs posterior for harsh touch; sets the expected response."),
          ("Design", "single / habituation series / independent animals; controls worm identity and cross-trial comparison."),
          ("Stimulus times or on-movie marks", "When each stimulus was delivered."),
          ("Optional blackout window", "Excludes pick-in-view frames from scoring."),
          ("FPS and scale", "Calibrate velocity and amplitude.")],
         ["The single-worm tracker produces per-frame signed velocity and body curvature. Around each stimulus the tool splits time into before, response, and after windows and reports, in each, the mean crawling velocity, the head-bend amplitude (a robust 95th-minus-5th-percentile of the head-bend angle), a Tierpsy-style box-aspect-ratio quirkiness from the spine, and a centroid path tortuosity as a spine-independent fallback. Reversals are scored from the signed velocity, and each trial is classified as a reversal, a forward escape, a stop without reversal, or no response. The habituation analysis fits the decay of response probability across trials at the plate level."],
         ["plate_trial_series.csv (response fraction per trial per plate); reversal_window_metrics.csv (before/during/after velocity, head-bend amplitude, quirkiness, tortuosity, plus reversed and stopped-without-reversal flags); spontaneous_reversals.csv and a summary for the spontaneous mode; and the plate habituation fit."],
         ["A blank column reads NaN honestly when a window has too few valid frames or a missing input. Poor segmentation drops the quirkiness column but tortuosity still reports. A wrong scale or FPS rescales velocity and amplitude directly."],
         ["Use the blackout window when the pick lingers; use the spontaneous design when there is no stimulus; keep the plate as the unit of inference for population designs."])
    tool(doc, "Population tap response and habituation", "Experimental",
         "Score how many worms in a population respond to a plate tap, and how strongly, from centroid trajectories.",
         "You have a population recording with a plate tap and want the fraction of animals responding, the response strength, and habituation across taps, without needing per-worm spines.",
         ["The tap moves the whole field of view; its intensity is read from the size of that global-motion artifact, with a duration and an inter-tap frequency.",
          "Each worm's centroid track is split before and after each tap and paired, so every animal is its own control.",
          "Response is classified per worm as a change in speed and/or direction; it is centroid-based, so no clean spine is required.",
          "The plate is the unit of replication."],
         [("Population tracks CSV", "From the population tracker (track_id, frame, x, y)."),
          ("Movie", "For the global-motion tap signal."),
          ("Speed and direction thresholds", "Define what counts as a response."),
          ("Before and after windows", "The paired comparison windows."),
          ("FPS and scale", "Calibrate the trajectories.")],
         ["The tool measures per-frame global motion (mean absolute frame-to-frame difference), detects taps as runs above a robust threshold with an intensity, duration, and frequency, and for each worm and tap compares before-versus-after centroid speed and heading. It reports per-worm responder flags and the population response fraction per tap."],
         ["taps.csv (onset, intensity, duration, interval); per_worm_tap_response.csv (before/after speed and heading, responded flags); population_tap_summary.json (fraction responding, split by speed vs direction, and mean response strength, per tap)."],
         ["No tap is detected if the FPS is wrong or the artifact is weak. Very short tracks are marked non-trackable and excluded from the denominator rather than scored as non-responders."],
         ["Run the population tracker first to produce the tracks CSV; being centroid-only, it works when spines are unavailable."])


def utilities(doc):
    doc.add_heading("Part VIII - Acquisition and utilities", level=1)
    tool(doc, "Sample planner: how many more?", "Ready",
         "Decide whether you have enough replicates and, if not, how many more, using the honest test for your data.",
         "You have group values (or a module's plate-level export) and want the current power and the additional replicate units needed to reach a target.",
         ["The plate (or the chosen replicate unit) is n; worm-level power is refused for population assays as pseudoreplication.",
          "The data drives the test: it checks for outliers, normality (Shapiro-Wilk), and equal variance (Levene), then forks to Welch's t, Mann-Whitney, Welch ANOVA, or Kruskal-Wallis.",
          "It runs offline in the browser and no data leaves the machine."],
         [("Replicate unit", "Plate / well / worm; defines what counts as n."),
          ("Groups", "Two or three-plus groups, pasted or loaded from a CSV."),
          ("Power target and alpha", "Confidence to catch a real effect and the false-alarm rate."),
          ("CSV group and value columns", "When loading a module export, which column splits conditions and which is the metric.")],
         ["From the group values the planner computes the standardized effect (the gap divided by the pooled scatter), the current power, and the n per group required for the target power, and draws the power-versus-n curve so diminishing returns are visible. The statistical core is validated against SciPy."],
         ["An on-screen verdict (you are there / almost / keep going / no effect yet), the current n and power, the target n, and how many more units to run, with every data check shown alongside its reasoning."],
         ["Under about six units the variance estimate is itself noisy, so the number is a target to move toward and re-check as data arrive. More replicates will not manufacture an effect that is genuinely small."],
         ["Load a plate-level export (for example plate_trial_series.csv from an assay) and choose the group and value columns to plan directly from a module's output."])
    tool(doc, "Probe a movie", "Ready",
         "Determine what a source file or image sequence contains and whether its technical properties suit behavioral or fluorescence analysis.",
         "The format, frame count, dimensions, codec, bit depth, or metadata are uncertain.",
         ["A preflight tool prevents silent analysis of the wrong temporal or intensity representation."],
         [("Selected video/stack/folder", "Defines the source to inspect.")],
         ["The probe reads the container or image sequence and reports accessible acquisition/format properties without altering the source."],
         ["Human-readable technical report used to plan conversion or analysis."],
         ["Missing container metadata; variable frame rate; codec libraries; a file reporting nominal FPS that differs from effective timestamps."],
         ["Verify critical acquisition settings against microscope/camera records, not the container alone."])

    tool(doc, "Convert for Fiji", "Ready",
         "Create a clean Fiji-readable TIFF stack from a movie or image source.",
         "Fiji cannot open the source reliably or a standardized stack is needed.",
         ["Conversion is treated as a derivative, not a replacement for raw data.",
          "Optional trimming, frame stepping, and spatial scaling are exposed because file size and temporal/spatial resolution trade off against biological resolvability."],
         [("Start/stop frame", "Trim to the relevant interval without changing source."),
          ("Frame step", "Reduce temporal density; this also reduces resolvable frequency."),
          ("Spatial scale percent", "Reduce size; this also reduces anatomical detail."),
          ("Output destination", "Keeps the derivative separate from raw data.")],
         ["The converter decodes frames, applies requested sampling/scaling, and writes a clean TIFF stack."],
         ["Converted stack and conversion summary."],
         ["Loss of temporal resolution from stepping; loss of spatial detail from scaling; altered bit depth; codec decode differences; large output files."],
         ["Never downsample before deciding the fastest biological event and smallest structure that must remain resolvable."])

    tool(doc, "Install AGVGLab Fiji menu", "Ready utility",
         "Install the lab's Fiji tools into a consistent Plugins > AGVGLab menu.",
         "A lab computer needs the Fiji integrations installed or refreshed.",
         ["Centralized menu installation reduces student dependence on dragging scripts and improves version consistency."],
         [("Fiji installation location", "Targets the correct local Fiji instance.")],
         ["The installer copies/registers the supplied Fiji tools into the expected menu structure."],
         ["Installed menu entries; no biological measurements."],
         ["Multiple Fiji installations; permissions; stale older plugins; Java/ImageJ version differences."],
         ["After an update, confirm which Fiji executable the Hub uses and archive the plugin version with results."])


def closing(doc):
    doc.add_heading("Part IX - Validation, statistics, troubleshooting, and maintenance", level=1)
    doc.add_heading("Validation levels", level=2)
    add_matrix(doc, ["Level", "Meaning", "Permitted claim"], [
        ("Computational regression", "Known synthetic or fixture input produces expected software output", "The implementation behaves as tested"),
        ("Technical validation", "Agreement with manual scoring or a reference method across representative recordings", "The measurement is technically credible under tested conditions"),
        ("Biological validation", "Expected and novel phenotypes are evaluated across animals, conditions, and confounds", "The readout supports biological inference in the validated domain"),
        ("Publication use", "Methods, QC, exclusions, replication, and version are fully reported", "Results are reproducible and interpretable by others"),
    ], [1700, 4100, 3560])
    add_callout(doc, "Experimental label",
                "An experimental tool may be useful and regression-tested while still lacking broad technical or biological validation. "
                "The label is a scientific boundary, not a cosmetic warning.")
    doc.add_heading("Correct units of inference", level=2)
    add_bullets(doc, [
        "One worm remains one animal even when it contributes many frames, segments, transients, or lawn encounters.",
        "One plate remains one independent orientation assay even when it contains many worms.",
        "Repeated events require paired, repeated-measures, hierarchical, or mixed-effects analysis.",
        "Technical repeats and reanalyses of one recording do not increase biological N.",
        "Report exclusions and uncertain fractions by condition; QC itself can reveal a phenotype or acquisition bias."])
    doc.add_heading("General troubleshooting sequence", level=2)
    add_numbered(doc, [
        "Confirm the source: correct file, frame order, channels, FPS, scale, exposure, and bit depth.",
        "Inspect raw frames before changing thresholds.",
        "Check whether the biological feature is visibly resolvable.",
        "Inspect segmentation/midline/track/ROI overlays.",
        "Identify whether failure is detection, identity, anatomy, timing, calibration, or inference.",
        "Change one exposed variable at a time and save the changed metadata.",
        "Re-run a known control or fixture after modifying code or defaults.",
        "If the needed evidence is absent, reacquire rather than tuning indefinitely."])
    doc.add_heading("Minimum methods reporting", level=2)
    add_bullets(doc, [
        "Lab Tools installer/manual version and tool name/status.",
        "Biological preparation, strain, age, condition, temperature, assay geometry, and experimental unit.",
        "Microscope/camera, objective, FPS, scale, exposure, bit depth, compression, and recording duration.",
        "All non-default thresholds, windows, buffers, modes, ROIs, anatomical orientation, and exclusion rules.",
        "Automated candidate count, reviewed accepted/rejected/uncertain count, and QC retention.",
        "Definition of every primary metric and the statistical model at the correct unit of inference."])
    doc.add_heading("Maintaining this manual as tools evolve", level=2)
    add_numbered(doc, [
        "Give every release a version and date.",
        "Update the Hub registry description and entry requirements.",
        "Update the tool chapter's biological aim and 'use when' boundary.",
        "Record any algorithmic decision that changes what can be detected or compared.",
        "Document every new user-adjustable variable and why it cannot be fixed universally.",
        "Update outputs, provenance fields, known failure modes, validation status, and regression tests.",
        "Render and visually inspect the revised manual before distribution.",
    ])
    doc.add_heading("Glossary", level=2)
    add_matrix(doc, ["Term", "Meaning in this ecosystem"], [
        ("Candidate", "An automated proposal awaiting QC or human review"),
        ("Reviewed result", "A candidate set altered or accepted by an identified review workflow"),
        ("ROI", "Saved spatial geometry that defines or filters the biological field"),
        ("Provenance", "Acquisition, settings, software version, and manual decisions needed to reconstruct a result"),
        ("Coverage", "The fraction/duration of data that supports a measurement"),
        ("Tracklet stitch", "An auditable inferred continuation across a short detection gap"),
        ("Uncertain", "Evidence is insufficient or conflicting; not a fourth biological behavior by default"),
        ("Experimental tool", "Implemented and testable, but not yet broadly validated for unattended biological inference"),
    ], [2200, 7160])
    doc.add_heading("Source-of-truth note", level=2)
    doc.add_paragraph(
        f"This manual describes the staged WINK Lab Hub v{VERSION} registry and implementation snapshot. "
        "Where a user interface, exported metadata file, or current source code differs from this manual, "
        "pause the analysis and reconcile the software/manual versions before interpreting results.")


def build():
    doc = Document()
    configure(doc)
    title_page(doc)
    front_matter(doc)
    shared_principles(doc)
    locomotion(doc)
    rhythmic(doc)
    physiology(doc)
    morphology(doc)
    mechanosensation(doc)
    utilities(doc)
    closing(doc)
    core = doc.core_properties
    core.title = "AGVG Lab Tools: Biology-First User and Methods Manual"
    core.subject = "Ultimate manual for the Vidal-Gadea Lab C. elegans analysis ecosystem"
    core.author = "Vidal-Gadea Lab"
    core.keywords = "C. elegans, behavior, kinematics, calcium, morphology, Lab Hub"
    core.comments = f"Living manual; software snapshot v{VERSION}"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
